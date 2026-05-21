"""Shared helpers for IEEE paper generation."""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

CHARTS = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'paper_charts')

def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    s = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color_hex
    })
    shading.append(s)

def add_heading_ieee(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.size = Pt(12 if level == 1 else 11)
    return h

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Pt(18)
    for r in p.runs:
        r.font.name = 'Times New Roman'
    return p

def add_body_no_indent(doc, text):
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_figure(doc, img_name, caption):
    path = os.path.join(CHARTS, img_name)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(path, width=Inches(5.0))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.name = 'Times New Roman'

def make_table(doc, headers, rows, caption=None):
    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(9)
            r.font.italic = True
            r.font.name = 'Times New Roman'
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        for par in cell.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
        set_cell_shading(cell, '1F4E79')
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for par in cell.paragraphs:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in par.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
            if ri % 2 == 1:
                set_cell_shading(cell, 'D6E4F0')
    doc.add_paragraph()
    return tbl
